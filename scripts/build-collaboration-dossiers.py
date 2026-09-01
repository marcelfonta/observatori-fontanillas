from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "patrocini" / "dossiers"
LOGO = ROOT / "assets" / "images" / "observatori-fontanillas-avatar-v21.png"

GREEN = "286D55"
GREEN_DARK = "123F31"
GREEN_LIGHT = "E8F3EE"
MINT = "71CFA0"
INK = "173129"
MUTED = "5D6F68"
LINE = "C9DAD3"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=130, start=150, bottom=130, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def font(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    font(run, 8.5, MUTED)


def configure_document(doc, labels):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    style_tokens = {
        "Title": (30, GREEN_DARK, 0, 8),
        "Subtitle": (14, MUTED, 0, 16),
        "Heading 1": (19, GREEN_DARK, 15, 8),
        "Heading 2": (13, GREEN, 10, 4),
        "Heading 3": (11, GREEN_DARK, 7, 3),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Word's built-in Title style can inherit a coloured bottom border from
    # the document theme. The dossier cover uses a clean, borderless title.
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.86))
    set_table_geometry(table, [4680, 4680], indent=0)
    table.cell(0, 0).text = labels["brand"]
    table.cell(0, 1).text = labels["header"]
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in table.row_cells(0):
        set_cell_margins(cell, top=0, start=0, bottom=60, end=0)
        for run in cell.paragraphs[0].runs:
            font(run, 8.2, GREEN if cell is table.cell(0, 0) else MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    font(p.add_run("meteo.fontanillas.cat"), 8.5, GREEN, bold=True)
    font(p.add_run("  ·  "), 8.5, MUTED)
    font(p.add_run(labels["footer"]), 8.5, MUTED)
    add_page_field(footer.add_paragraph())


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text.upper())
    font(run, 8.5, GREEN, bold=True)
    run.font.letter_spacing = Pt(1.1)
    return p


def add_lead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.28
    font(p.add_run(text), 12.2, GREEN_DARK, bold=True)
    return p


def add_card_grid(doc, cards):
    table = doc.add_table(rows=1, cols=len(cards))
    widths = [9360 // len(cards)] * len(cards)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for index, (title, body) in enumerate(cards):
        cell = table.cell(0, index)
        set_cell_shading(cell, GREEN_LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        font(p.add_run(title), 11, GREEN_DARK, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.16
        font(p2.add_run(body), 9.3, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_label_paragraph(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run(label + "  "), 10.7, GREEN_DARK, bold=True)
    font(p.add_run(body), 10.5, INK)


def add_two_col_rows(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2760, 6600])
    table.cell(0, 0).text = rows[0][0]
    table.cell(0, 1).text = rows[0][1]
    for label, body in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = body
    set_table_geometry(table, [2760, 6600])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if col_index == 0:
                set_cell_shading(cell, GREEN_LIGHT)
            for p in cell.paragraphs:
                for run in p.runs:
                    font(run, 9.5, GREEN_DARK if col_index == 0 else INK, bold=col_index == 0)
        if row_index == 0:
            set_repeat_table_header(row)
    return table


def new_page(doc, kicker, title, intro=None):
    doc.add_page_break()
    add_kicker(doc, kicker)
    doc.add_paragraph(title, style="Heading 1")
    if intro:
        add_lead(doc, intro)


def build(lang, filename):
    es = lang == "es"
    labels = {
        "brand": "OBSERVATORIO FONTANILLAS" if es else "FONTANILLAS WEATHER OBSERVATORY",
        "header": "Dossier de colaboración" if es else "Collaboration dossier",
        "footer": "Propuesta de material cedido" if es else "Equipment collaboration proposal",
    }
    doc = Document()
    configure_document(doc, labels)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.22))
    doc.inline_shapes[-1]._inline.docPr.set(
        "descr",
        "Logotipo del Observatorio Meteorológico Fontanillas" if es
        else "Fontanillas Weather Observatory logo",
    )
    add_kicker(doc, "Sant Celoni · Montseny · 2026")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = "Propuesta de colaboración técnica" if es else "Technical collaboration proposal"
    subtitle = "Observación meteorológica local, integración de datos y divulgación transparente" if es else "Local weather observation, data integration and transparent outreach"
    doc.add_paragraph(title, style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(subtitle, style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    font(p.add_run("meteo.fontanillas.cat"), 12, GREEN, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p2.add_run("meteo@fontanillas.cat"), 10.5, MUTED)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(36)
    font(p3.add_run("CAT · ES · EN · FR"), 9.5, GREEN_DARK, bold=True)

    if es:
        new_page(doc, "El proyecto", "Un observatorio local con vocación pública", "El Observatorio Fontanillas convierte datos meteorológicos reales de Sant Celoni en información comprensible, verificable y útil para la comunidad.")
        doc.add_paragraph("Qué reúne", style="Heading 2")
        add_card_grid(doc, [
            ("Observación", "Lecturas en directo, histórico meteorológico y contexto de los sensores."),
            ("Previsión", "Pronósticos diferenciados de la observación, con verificación posterior."),
            ("Divulgación", "Recursos educativos y piezas meteorológicas adaptadas a web y redes."),
        ])
        doc.add_paragraph("Una base técnica ya operativa", style="Heading 2")
        add_label_paragraph(doc, "Portal multilingüe.", "Contenido público en catalán, castellano, inglés y francés, con buscador de municipios y experiencia móvil.")
        add_label_paragraph(doc, "Datos y automatización.", "Web en tiempo real, archivo histórico, avisos oficiales separados, verificación de previsiones y formatos sociales automatizados.")
        add_label_paragraph(doc, "Entorno de campo.", "Sant Celoni se encuentra entre la llanura vallesana y el macizo del Montseny, un contexto útil para documentar lluvia, viento, radiación y cambios térmicos locales.")
        doc.add_paragraph("Principio editorial", style="Heading 2")
        p = doc.add_paragraph()
        set_para = p.paragraph_format
        set_para.left_indent = Inches(.25)
        set_para.right_indent = Inches(.25)
        font(p.add_run("Los datos observados, las previsiones, los avisos oficiales y la interpretación se presentan siempre como categorías distintas."), 11.2, GREEN_DARK, bold=True)

        new_page(doc, "La colaboración", "Material útil, resultados públicos", "La propuesta no requiere una aportación económica. Buscamos equipos cedidos o prestados y apoyo técnico para construir un caso real y bien documentado.")
        add_two_col_rows(doc, [
            ("Aportación del colaborador", "Estación, gateway, sensor o acceso técnico acordado; documentación y condiciones de uso aplicables."),
            ("Trabajo del Observatorio", "Revisión de compatibilidad, instalación, integración, seguimiento, documentación y publicación."),
            ("Resultado público", "Página o artículo técnico, explicación educativa, atribución visible y comunicación en los canales del proyecto."),
            ("Duración", "Se acuerda según el equipo: prueba puntual, préstamo temporal o incorporación estable."),
        ])
        doc.add_paragraph("Entregables posibles", style="Heading 2")
        add_card_grid(doc, [
            ("Instalación", "Criterios de ubicación, conectividad y mantenimiento explicados de forma comprensible."),
            ("Integración", "Flujo de datos y API documentado, sin publicar credenciales ni saltarse licencias."),
            ("Caso práctico", "Resultados, incidencias y límites publicados con una atribución clara del material."),
        ])
        doc.add_paragraph("Redes y publicaciones periódicas", style="Heading 2")
        add_label_paragraph(doc, "Canales.", "Instagram, Facebook, TikTok, YouTube y X, con acceso al resto de comunidades desde el directorio oficial de redes del proyecto.")
        add_label_paragraph(doc, "Continuidad.", "Una presentación inicial, actualizaciones periódicas acordadas durante la colaboración y una pieza final con resultados, con una cadencia que no sature a la audiencia.")
        add_label_paragraph(doc, "Identificación.", "Cada pieza indicará claramente si existe material cedido, préstamo o patrocinio y respetará las normas publicitarias de cada plataforma.")
        doc.add_paragraph("Sin promesas artificiales", style="Heading 2")
        add_label_paragraph(doc, "No garantizamos alcance.", "No se ofrecen cifras de audiencia que no estén verificadas ni se compra interacción.")
        add_label_paragraph(doc, "No garantizamos una opinión positiva.", "La colaboración permite probar y explicar; las conclusiones permanecen independientes.")

        new_page(doc, "Aplicaciones", "Equipos que pueden aportar valor", "El alcance se adapta al producto y a sus condiciones técnicas. Priorizamos una integración pequeña, sostenible y medible.")
        add_two_col_rows(doc, [
            ("Estación completa", "Temperatura, humedad, presión, precipitación y viento para una evaluación integral."),
            ("Pluviómetro", "Comparación de episodios, resolución, mantenimiento y continuidad del dato."),
            ("Viento", "Seguimiento de rachas, dirección y decisiones de ubicación del sensor."),
            ("Radiación y UV", "Contexto solar, confort y recursos educativos sobre exposición."),
            ("Aire y ambiente", "Partículas, CO₂ u otras variables, claramente separadas de los avisos oficiales."),
            ("Suelo o vegetación", "Humedad, temperatura o mojado foliar para aplicaciones ambientales y educativas."),
            ("Rayos", "Detección y divulgación del fenómeno sin sustituir radares ni alertas oficiales."),
        ])
        doc.add_paragraph("Compatibilidad antes de instalar", style="Heading 2")
        add_label_paragraph(doc, "Conectividad.", "Se revisan Wi‑Fi, radio, alimentación, alcance y recuperación frente a fallos.")
        add_label_paragraph(doc, "Datos.", "Se documentan propiedad, frecuencia, retención, API, exportación y atribución.")
        add_label_paragraph(doc, "Mantenimiento.", "Se acuerdan limpieza, calibración disponible, consumibles, garantía y devolución si es un préstamo.")

        new_page(doc, "Garantías", "Transparencia y responsabilidad", "Una colaboración solo tiene sentido si mejora la observación sin confundir publicidad, datos y conclusiones.")
        add_card_grid(doc, [
            ("Divulgación visible", "Toda cesión, préstamo o acceso técnico se identifica en la página correspondiente."),
            ("Independencia", "La marca no controla resultados, redacción ni conclusiones meteorológicas."),
            ("Trazabilidad", "Se mantienen visibles la fuente, el tipo de dato, la escala y las limitaciones."),
        ])
        doc.add_paragraph("Compromisos del Observatorio", style="Heading 2")
        add_label_paragraph(doc, "Rigor.", "No inventar datos, sensores, estados ni beneficios del producto.")
        add_label_paragraph(doc, "Seguridad.", "No exponer claves, cuentas, ubicaciones sensibles ni accesos administrativos.")
        add_label_paragraph(doc, "Licencias.", "Usar datos, API, imágenes y logotipos solo con permiso y dentro de las condiciones aplicables.")
        add_label_paragraph(doc, "Avisos.", "Meteocat, AEMET, Protección Civil y 112 prevalecen siempre sobre cualquier resumen propio.")
        add_label_paragraph(doc, "Privacidad.", "No ceder datos personales de usuarios ni convertir la colaboración en seguimiento publicitario.")
        doc.add_paragraph("Qué no incluye", style="Heading 2")
        add_label_paragraph(doc, "Publicidad encubierta.", "La relación no se ocultará ni se presentará como una recomendación espontánea.")
        add_label_paragraph(doc, "Exclusividad automática.", "Cualquier exclusividad tendría que negociarse expresamente y nunca afectaría a la independencia de los datos.")

        new_page(doc, "Siguiente paso", "Una prueba pequeña y bien definida", "La mejor forma de empezar es acordar un equipo, una pregunta de evaluación y un periodo de prueba razonable.")
        add_two_col_rows(doc, [
            ("1. Propuesta", "Producto o sensor, modalidad de cesión y objetivo de la colaboración."),
            ("2. Revisión", "Compatibilidad técnica, ubicación, datos, licencias, atribución y mantenimiento."),
            ("3. Acuerdo", "Entregables, calendario, devolución si procede y declaración pública de la relación."),
            ("4. Ejecución", "Instalación, validación, integración y seguimiento con incidencias documentadas."),
            ("5. Publicación", "Caso práctico, recursos educativos y conclusiones independientes."),
        ])
        doc.add_paragraph("Contacto", style="Heading 2")
        add_label_paragraph(doc, "Responsable.", "Marcel Fontanillas")
        add_label_paragraph(doc, "Correo.", "meteo@fontanillas.cat")
        add_label_paragraph(doc, "Web.", "https://meteo.fontanillas.cat")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run("OBSERVACIÓN LOCAL · DATOS VERIFICABLES · DIVULGACIÓN ABIERTA"), 9.5, GREEN, bold=True)
    else:
        new_page(doc, "The project", "A public-minded local observatory", "The Fontanillas Weather Observatory turns real weather observations from Sant Celoni into clear, verifiable and useful information for the community.")
        doc.add_paragraph("What it brings together", style="Heading 2")
        add_card_grid(doc, [
            ("Observation", "Live readings, a weather archive and clear sensor context."),
            ("Forecasting", "Forecasts kept separate from observations and verified afterwards."),
            ("Outreach", "Educational resources and weather content for web and social channels."),
        ])
        doc.add_paragraph("An operational technical foundation", style="Heading 2")
        add_label_paragraph(doc, "Multilingual portal.", "Public content in Catalan, Spanish, English and French, with municipality search and a mobile-first experience.")
        add_label_paragraph(doc, "Data and automation.", "Live website, historical archive, separate official alerts, forecast verification and automated social formats.")
        add_label_paragraph(doc, "Field setting.", "Sant Celoni lies between the Vallès plain and the Montseny massif, a useful setting for documenting local rain, wind, radiation and temperature changes.")
        doc.add_paragraph("Editorial principle", style="Heading 2")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(.25)
        p.paragraph_format.right_indent = Inches(.25)
        font(p.add_run("Observed data, forecasts, official alerts and interpretation are always presented as distinct categories."), 11.2, GREEN_DARK, bold=True)

        new_page(doc, "The collaboration", "Useful equipment, public results", "No financial contribution is required. We are seeking donated or loaned equipment and technical support to build a genuine, well-documented field case.")
        add_two_col_rows(doc, [
            ("Partner contribution", "An agreed station, gateway, sensor or technical access, plus applicable documentation and terms."),
            ("Observatory work", "Compatibility review, installation, integration, monitoring, documentation and publication."),
            ("Public outcome", "A technical page or article, educational explanation, visible attribution and communication through project channels."),
            ("Duration", "Defined for the equipment: a short test, temporary loan or stable installation."),
        ])
        doc.add_paragraph("Possible deliverables", style="Heading 2")
        add_card_grid(doc, [
            ("Installation", "Siting, connectivity and maintenance criteria explained in practical terms."),
            ("Integration", "A documented data and API workflow with no exposed credentials or licence bypasses."),
            ("Case study", "Results, issues and limitations published with clear equipment attribution."),
        ])
        doc.add_paragraph("Social channels and periodic posts", style="Heading 2")
        add_label_paragraph(doc, "Channels.", "Instagram, Facebook, TikTok, YouTube and X, with access to the other communities through the project’s official social directory.")
        add_label_paragraph(doc, "Continuity.", "An initial announcement, agreed periodic updates during the collaboration and a final results piece, at a cadence designed not to overwhelm the audience.")
        add_label_paragraph(doc, "Disclosure.", "Every item will clearly identify donated or loaned equipment or sponsorship and follow each platform’s advertising rules.")
        doc.add_paragraph("No artificial promises", style="Heading 2")
        add_label_paragraph(doc, "No reach guarantee.", "We do not offer unverified audience figures or purchase engagement.")
        add_label_paragraph(doc, "No positive-review guarantee.", "The collaboration enables testing and explanation; conclusions remain independent.")

        new_page(doc, "Applications", "Equipment that can add value", "The scope is adapted to the product and its technical terms. We favour a small, sustainable and measurable integration.")
        add_two_col_rows(doc, [
            ("Complete station", "Temperature, humidity, pressure, precipitation and wind for an integrated assessment."),
            ("Rain gauge", "Event comparison, resolution, maintenance and data continuity."),
            ("Wind", "Gust and direction monitoring, including sensor-siting decisions."),
            ("Solar and UV", "Solar context, comfort and educational resources about exposure."),
            ("Air and environment", "Particulate matter, CO₂ or other variables, clearly separated from official alerts."),
            ("Soil or vegetation", "Moisture, temperature or leaf wetness for environmental and educational uses."),
            ("Lightning", "Detection and public explanation without replacing official radar or alerts."),
        ])
        doc.add_paragraph("Compatibility before installation", style="Heading 2")
        add_label_paragraph(doc, "Connectivity.", "Wi‑Fi, radio, power, range and failure recovery are reviewed first.")
        add_label_paragraph(doc, "Data.", "Ownership, frequency, retention, API, export and attribution are documented.")
        add_label_paragraph(doc, "Maintenance.", "Cleaning, available calibration, consumables, warranty and return conditions are agreed for loans.")

        new_page(doc, "Safeguards", "Transparency and responsibility", "A collaboration is worthwhile only if it improves observation without confusing advertising, data and conclusions.")
        add_card_grid(doc, [
            ("Visible disclosure", "Every donation, loan or technical access arrangement is identified on the relevant page."),
            ("Independence", "The brand does not control results, wording or meteorological conclusions."),
            ("Traceability", "Source, data type, scale and limitations remain visible."),
        ])
        doc.add_paragraph("Observatory commitments", style="Heading 2")
        add_label_paragraph(doc, "Accuracy.", "Never invent data, sensors, states or product benefits.")
        add_label_paragraph(doc, "Security.", "Never expose keys, accounts, sensitive locations or administrative access.")
        add_label_paragraph(doc, "Licensing.", "Use data, APIs, images and logos only with permission and under the applicable terms.")
        add_label_paragraph(doc, "Alerts.", "Meteocat, AEMET, Civil Protection and emergency services always take precedence over project summaries.")
        add_label_paragraph(doc, "Privacy.", "Do not share personal user data or turn the collaboration into advertising tracking.")
        doc.add_paragraph("Not included", style="Heading 2")
        add_label_paragraph(doc, "Hidden advertising.", "The relationship will not be concealed or presented as an unsolicited recommendation.")
        add_label_paragraph(doc, "Automatic exclusivity.", "Any exclusivity would require a separate agreement and could never affect data independence.")

        new_page(doc, "Next step", "A small, well-defined trial", "The best starting point is one agreed device, one evaluation question and a reasonable trial period.")
        add_two_col_rows(doc, [
            ("1. Proposal", "Product or sensor, equipment arrangement and collaboration objective."),
            ("2. Review", "Technical compatibility, siting, data, licences, attribution and maintenance."),
            ("3. Agreement", "Deliverables, schedule, return terms where relevant and public disclosure."),
            ("4. Delivery", "Installation, validation, integration and monitoring with documented issues."),
            ("5. Publication", "Case study, educational resources and independent conclusions."),
        ])
        doc.add_paragraph("Contact", style="Heading 2")
        add_label_paragraph(doc, "Project lead.", "Marcel Fontanillas")
        add_label_paragraph(doc, "Email.", "meteo@fontanillas.cat")
        add_label_paragraph(doc, "Website.", "https://meteo.fontanillas.cat")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run("LOCAL OBSERVATION · VERIFIABLE DATA · OPEN OUTREACH"), 9.5, GREEN, bold=True)

    doc.core_properties.title = title
    doc.core_properties.subject = labels["header"]
    doc.core_properties.author = "Observatori Meteorològic Fontanillas"
    doc.core_properties.keywords = "weather, observatory, collaboration, equipment, Sant Celoni, Montseny"
    OUT.mkdir(parents=True, exist_ok=True)
    output = OUT / filename
    doc.save(output)
    return output


if __name__ == "__main__":
    print(build("es", "Dossier-colaboracion-Observatorio-Fontanillas-ES.docx"))
    print(build("en", "Collaboration-dossier-Fontanillas-Weather-Observatory-EN.docx"))
